#!/usr/bin/env python3
"""Generate a Word document for the CI + US Expansion forecast assumptions."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# -- Style setup --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9.5)

    return table

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Harness Give\nForecast Assumptions")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Scenario C: With Charity Intelligence + US Expansion")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x4A, 0x6C, 0x9E)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f"Prepared: {datetime.date.today().strftime('%B %d, %Y')}\nConfidential — For Internal Use Only")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading("Contents", level=1)
toc_items = [
    "1. Scenario Overview",
    "2. Canadian Traffic Model (With CI)",
    "3. CI Traffic Analysis & Projections",
    "4. Canadian Registration Funnel",
    "5. US Market Entry & Bootstrap Mechanism",
    "6. US Traffic Model",
    "7. US Registration Funnel",
    "8. Combined User Growth (CA + US)",
    "9. Revenue Model",
    "10. Donation Seasonality",
    "11. Infrastructure Cost Model",
    "12. Marketing & User Acquisition",
    "13. Organizational Costs — All Four Entities",
    "14. Key Modeling Decisions",
    "15. Known Gaps",
    "Sources & References",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. SCENARIO OVERVIEW
# ============================================================
doc.add_heading("1. Scenario Overview", level=1)

doc.add_paragraph(
    "Scenario C models a Canadian launch with the Charity Intelligence (CI) partnership "
    "(identical to Scenario A) that bootstraps US market entry four months later. The proven "
    "Canadian platform and \"Charity Intelligence\" trust signal accelerate US growth versus "
    "a cold-market entry."
)

doc.add_paragraph(
    "Timeline: Canada launches April 2026 (FY2027 start). The US enters August 2026 "
    "(4 months later). FY2027 US figures reflect 8 months of activity."
)

doc.add_heading("Four-Entity Group Structure", level=2)
add_table(doc,
    ["Entity", "Jurisdiction", "Type", "Standards"],
    [
        ["Harness Exchange", "Canada (CBCA / Ontario)", "For-Profit", "ASPE"],
        ["Harness Good", "Canada (CNCA)", "Not-for-Profit", "ASNFPO"],
        ["Harness Exchange US", "Delaware C-Corp", "For-Profit", "US GAAP"],
        ["Harness Good US Foundation", "Delaware 501(c)(3)", "Not-for-Profit", "ASC 958"],
    ],
)

doc.add_paragraph()

# ============================================================
# 2. CANADIAN TRAFFIC MODEL
# ============================================================
doc.add_heading("2. Canadian Traffic Model (With CI)", level=1)

doc.add_paragraph(
    "The CI partnership provides the dominant traffic channel for the Canadian market. "
    "CI referral traffic is supplemented by organic search and paid media."
)

add_table(doc,
    ["Year", "CI Referral", "Organic Search", "Paid Media", "Total Visitors"],
    [
        ["FY2027", "90,000", "22,000", "18,000", "130,000"],
        ["FY2028", "100,000", "40,000", "35,000", "175,000"],
        ["FY2029", "120,000", "50,000", "40,000", "210,000"],
        ["FY2030", "150,000", "60,000", "50,000", "260,000"],
        ["FY2031", "180,000", "75,000", "65,000", "320,000"],
    ],
)

doc.add_paragraph()
doc.add_paragraph(
    "CI referral traffic represents 69% of Year 1 visitors at zero incremental cost — "
    "the primary value proposition of the partnership. Organic search starts at 22K "
    "(SEO takes 6–12 months for meaningful results) and grows to 75K by Year 5. Paid media "
    "scales from Year 3 onward to compensate for CI's constrained ceiling."
)

# ============================================================
# 3. CI TRAFFIC ANALYSIS
# ============================================================
doc.add_heading("3. CI Traffic Analysis & Projections", level=1)

doc.add_heading("Historical CI Traffic (Server-Side + GA4)", level=2)
doc.add_paragraph(
    "CI's traffic has been declining since its 2022 peak of 594,769 visits "
    "(533K unique visitors). The decline accelerated through 2025."
)

add_table(doc,
    ["Year", "CI Visits", "CI Unique Visitors", "YoY Change (Visits)"],
    [
        ["2020", "570,611", "516,352", "+10.8%"],
        ["2021", "526,391", "478,648", "-7.7%"],
        ["2022", "594,769", "533,056", "+13.0% (Ukraine crisis peak)"],
        ["2023", "465,210", "400,097", "-21.8% (GA4 migration + structural)"],
        ["2024", "447,732", "363,364", "-3.8%"],
        ["2025*", "387,290", "312,760", "-13.5% (annualized ~422K/~341K)"],
    ],
)

doc.add_paragraph()
p = doc.add_paragraph("Source: ")
p.add_run("CI Webstats daily log (server-side, 2009–2025) and GA4 (2023–2025). "
           "CI Web Traffic Analysis, prepared April 2026 (internal document).").italic = True

doc.add_heading("Decline Drivers", level=2)
doc.add_paragraph(
    "Increased competition from AI-generated charity information (ChatGPT drove 3,036+ "
    "referral pageviews in 2025), Google SERP features reducing click-through rates, and "
    "aging Joomla CMS hurting SEO performance. The Harness platform migration (modern CMS, "
    "mobile-first, structured data) is expected to stabilize and partially reverse this trend."
)

doc.add_heading("CI Unique Visitor Projections", level=2)
add_table(doc,
    ["Year", "CI Unique Visitors", "Trajectory Assumption"],
    [
        ["FY2027", "305,000", "Continued decline slowing; migration in progress"],
        ["FY2028", "295,000", "Trough year; migration complete, SEO recovery begins"],
        ["FY2029", "315,000", "Recovery: modern platform, improved mobile, structured data"],
        ["FY2030", "340,000", "Growth: AI readiness, French content, enhanced E-E-A-T"],
        ["FY2031", "365,000", "Sustained growth; unlikely to return to 2022 peak (533K)"],
    ],
)

doc.add_heading("Referral Capture Rate", level=2)
doc.add_paragraph(
    "The capture rate is the percentage of CI unique visitors who click through to Give. "
    "Year 1 starts at 30% (simple cross-links, new/unfamiliar product) and grows to 49% "
    "by Year 5 as integration deepens. Capture above 50% is unlikely — CI's own data shows "
    "only 77% of pageviews are charity-related, and not all visitors have giving intent."
)

add_table(doc,
    ["Year", "CI Unique Visitors", "Capture Rate", "CI Referral", "Integration Level"],
    [
        ["FY2027", "305,000", "30%", "90,000", "Cross-links on charity profile pages"],
        ["FY2028", "295,000", "34%", "100,000", "Dedicated Give section in search results"],
        ["FY2029", "315,000", "38%", "120,000", "Embedded donation widgets on CI pages"],
        ["FY2030", "340,000", "44%", "150,000", "Co-branded giving flows, email cross-promo"],
        ["FY2031", "365,000", "49%", "180,000", "Deep integration: CI as primary Give funnel"],
    ],
)

doc.add_paragraph()
p = doc.add_paragraph("Source: ")
p.add_run(
    "Cross-site referral rates for partner integrations typically range 15–40% (loosely integrated) "
    "to 40–60% (deeply embedded). Partnerize State of the Nation 2024. CI-specific: 93% of pageviews "
    "and 77% of organic search clicks go to charity-related content (CI Web Traffic Analysis, "
    "Google Search Console Dec 2024–Nov 2025, internal)."
).italic = True

# ============================================================
# 4. CANADIAN REGISTRATION FUNNEL
# ============================================================
doc.add_heading("4. Canadian Registration Funnel", level=1)

add_table(doc,
    ["Year", "Total Visitors", "Registered", "Conversion"],
    [
        ["FY2027", "130,000", "1,300", "1.0%"],
        ["FY2028", "175,000", "5,000", "2.9%"],
        ["FY2029", "210,000", "14,000", "6.7%"],
        ["FY2030", "260,000", "30,000", "11.5%"],
        ["FY2031", "320,000", "55,000", "17.2%"],
    ],
)

doc.add_paragraph()
doc.add_paragraph(
    "Registration rate starts at 1.0% (cold traffic, new product) and improves to 17.2% "
    "by Year 5 as Give builds trust and CI integration becomes seamless. Industry average "
    "for nonprofit platforms is 1.8% visitor-to-registration (M+R Benchmarks 2024). "
    "Year 5 exceeds industry average because CI traffic is pre-qualified."
)

# ============================================================
# 5. US MARKET ENTRY
# ============================================================
doc.add_heading("5. US Market Entry & Bootstrap Mechanism", level=1)

doc.add_paragraph(
    "Canadian launch with CI bootstraps US market entry 4 months later. The proven "
    "Canadian platform and \"Charity Intelligence\" trust signal accelerate US growth "
    "versus a cold-market entry."
)

doc.add_heading("Bootstrap Advantages", level=2)
bullets = [
    "\"As featured by Charity Intelligence Canada\" trust signal in US marketing",
    "Cross-border charity referrals (many US 501(c)(3)s have Canadian affiliates)",
    "Proven platform with real user testimonials and charity ratings",
    "Press coverage from Canadian launch drives US tech/nonprofit media interest",
    "Existing compliance, security, and payment infrastructure to build on",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

# ============================================================
# 6. US TRAFFIC MODEL
# ============================================================
doc.add_heading("6. US Traffic Model", level=1)

add_table(doc,
    ["Year", "Organic Search", "Paid Media", "PR / Referral", "Total Visitors"],
    [
        ["FY2027 (8 mo)", "10,000", "40,000", "10,000", "60,000"],
        ["FY2028", "40,000", "130,000", "30,000", "200,000"],
        ["FY2029", "80,000", "270,000", "50,000", "400,000"],
        ["FY2030", "150,000", "420,000", "80,000", "650,000"],
        ["FY2031", "250,000", "580,000", "120,000", "950,000"],
    ],
)

doc.add_paragraph()
doc.add_paragraph(
    "The US market has no CI-equivalent traffic source. Paid media is the dominant channel, "
    "supplemented by organic search and PR/referral traffic leveraging the Canadian launch story."
)

# ============================================================
# 7. US REGISTRATION FUNNEL
# ============================================================
doc.add_heading("7. US Registration Funnel", level=1)

add_table(doc,
    ["Year", "Total Visitors", "Registered", "Conversion"],
    [
        ["FY2027 (8 mo)", "60,000", "1,200", "2.0%"],
        ["FY2028", "200,000", "12,000", "6.0%"],
        ["FY2029", "400,000", "40,000", "10.0%"],
        ["FY2030", "650,000", "95,000", "14.6%"],
        ["FY2031", "950,000", "200,000", "21.1%"],
    ],
)

doc.add_paragraph()
doc.add_paragraph(
    "US conversion rates start higher than Canada's cold-start (2.0% vs 1.0%) because "
    "the platform is already proven with Canadian testimonials and charity data. By Year 5, "
    "US registration rate (21.1%) exceeds Canada (17.2%) due to larger addressable market "
    "and stronger network effects at scale."
)

# ============================================================
# 8. COMBINED USER GROWTH
# ============================================================
doc.add_heading("8. Combined User Growth (CA + US)", level=1)

add_table(doc,
    ["Year", "CA Users", "US Users", "Combined", "Infrastructure Tier"],
    [
        ["FY2027", "1,300", "1,200", "2,500", "Startup \u2192 Growth (17%)"],
        ["FY2028", "5,000", "12,000", "17,000", "Growth \u2192 Scale (18%)"],
        ["FY2029", "14,000", "40,000", "54,000", "Scale \u2192 Enterprise (2%)"],
        ["FY2030", "30,000", "95,000", "125,000", "Scale \u2192 Enterprise (38%)"],
        ["FY2031", "55,000", "200,000", "255,000", "Enterprise"],
    ],
)

doc.add_paragraph()
doc.add_paragraph(
    "US overtakes Canada in users by Year 2, reflecting the ~10\u00d7 larger addressable market. "
    "Combined user growth pushes infrastructure into higher tiers faster than either market "
    "alone — reaching Enterprise tier by Year 5."
)

# ============================================================
# 9. REVENUE MODEL
# ============================================================
doc.add_heading("9. Revenue Model", level=1)

doc.add_heading("Canadian Transaction Fee Revenue", level=2)
add_table(doc,
    ["Parameter", "Year 1", "Year 5", "Source"],
    [
        ["Active donor %", "8\u201312%", "21%", "Industry avg: 10\u201315% (M+R 2024)"],
        ["Donations per donor/yr", "2.0\u20133.0", "4.8", "CanadaHelps avg: 3.2 txns/donor/yr"],
        ["Avg donation", "$320\u2013350", "$500", "CRA T1 2023: median $390, online avg $500"],
        ["Platform fee", "3.5%", "3.2%", "CanadaHelps: 4.0%; GoFundMe: 2.9%"],
    ],
)

doc.add_paragraph()
doc.add_heading("US Revenue Parameters", level=2)
add_table(doc,
    ["Parameter", "Year 1 (8 mo)", "Year 5", "Source"],
    [
        ["Active donor %", "8%", "19%", "Lower than CA; US market more competitive"],
        ["Donations per donor/yr", "2.0", "4.5", "US online avg lower frequency (Network for Good 2023)"],
        ["Avg donation (USD)", "$250", "$400", "US online avg $200\u2013250; trends up with maturity"],
        ["Platform fee", "3.5%", "3.2%", "Same trajectory as CA; competitive with US market"],
        ["Stripe fee (US)", "2.2% + $0.30", "2.2% + $0.30", "Stripe 501(c)(3) nonprofit discount"],
    ],
)

doc.add_paragraph()
doc.add_heading("Payment Processing (Stripe)", level=2)
doc.add_paragraph(
    "Stripe fees are modeled as a cost of revenue. Both CA and US use the nonprofit "
    "discount rate of 2.2% + $0.30 per transaction. Stripe offers verified 501(c)(3) / "
    "registered charity organizations this discounted rate (vs standard 2.9% + $0.30)."
)

doc.add_heading("Revenue Formula", level=3)
doc.add_paragraph(
    "Active Donors = Users \u00d7 Active Donor %\n"
    "Total Txns = Active Donors \u00d7 Donations/Donor\n"
    "Donation Volume = Total Txns \u00d7 Avg Donation\n"
    "Transaction Fee Revenue = Donation Volume \u00d7 Platform Fee %"
)

doc.add_heading("Net Revenue Per Transaction", level=3)
doc.add_paragraph(
    "Example — CA Year 1 ($350 avg donation, 3.5% fee):\n"
    "  Gross fee = $350 \u00d7 3.5% = $12.25\n"
    "  Stripe cost = $350 \u00d7 2.2% + $0.30 = $8.00\n"
    "  Net revenue = $4.25 (1.21% effective net margin)"
)
doc.add_paragraph(
    "Example — US Year 1 ($250 avg donation, 3.5% fee):\n"
    "  Gross fee = $250 \u00d7 3.5% = $8.75\n"
    "  Stripe cost = $250 \u00d7 2.2% + $0.30 = $5.80\n"
    "  Net revenue = $2.95 (1.18% effective net margin)"
)

doc.add_heading("Charity Partner SaaS Revenue", level=2)
add_table(doc,
    ["Year", "Partners", "Avg Annual Fee", "Total"],
    [
        ["FY2027", "1\u20135", "$600", "$600\u2013$3,000"],
        ["FY2028", "10\u201335", "$900", "$9,000\u2013$31,500"],
        ["FY2029", "40\u201380", "$1,200", "$48,000\u2013$96,000"],
        ["FY2030", "100\u2013150", "$1,500", "$150,000\u2013$225,000"],
        ["FY2031", "250\u2013280", "$1,800", "$450,000\u2013$504,000"],
    ],
)

# ============================================================
# 10. DONATION SEASONALITY
# ============================================================
doc.add_heading("10. Donation Seasonality", level=1)

doc.add_paragraph(
    "Donation revenue and Stripe processing fees follow a giving-season curve. "
    "SaaS revenue and infrastructure costs use uniform monthly distribution."
)

add_table(doc,
    ["Month", "Donation Weight", "Marketing Weight", "Rationale"],
    [
        ["Jan", "8%", "6%", "Post-holiday giving, New Year pledges"],
        ["Feb", "5%", "4%", "Baseline"],
        ["Mar", "6%", "5%", "Fiscal year-end giving"],
        ["Apr", "5%", "4%", "Baseline"],
        ["May", "5%", "4%", "Baseline"],
        ["Jun", "5%", "5%", "Mid-year campaigns"],
        ["Jul", "4%", "5%", "Summer low"],
        ["Aug", "4%", "6%", "Creative production, audience building"],
        ["Sep", "5%", "11%", "Campaign launch; early giving season"],
        ["Oct", "9%", "15%", "Giving season begins; Thanksgiving (CA)"],
        ["Nov", "13%", "18%", "GivingTuesday; pre-holiday campaigns"],
        ["Dec", "31%", "17%", "Peak holiday giving; ~30% of annual online donations"],
    ],
)

doc.add_paragraph()
p = doc.add_paragraph("Sources: ")
p.add_run(
    "Blackbaud Charitable Giving Report 2023 (30% of annual online giving in December); "
    "M+R Benchmarks 2024 (Nov + Dec = 26\u201336% of annual online revenue); "
    "CanadaHelps Giving Report 2023; Network for Good Digital Giving Trends 2023 "
    "(47% of online donations in Q4)."
).italic = True

# ============================================================
# 11. INFRASTRUCTURE COST MODEL
# ============================================================
doc.add_heading("11. Infrastructure Cost Model", level=1)

doc.add_paragraph(
    "Monthly per-tier estimates based on actual cloud provider pricing (Q1 2026). "
    "When a year falls between tiers, costs are linearly interpolated."
)

add_table(doc,
    ["Category", "Startup (1K)", "Growth (10K)", "Scale (50K)", "Enterprise (250K)"],
    [
        ["Compute & Hosting", "$35", "$85", "$400", "$1,050"],
        ["PostgreSQL + pgvector", "$18", "$100", "$375", "$1,400"],
        ["Redis Cache", "$5", "$35", "$150", "$300"],
        ["Kafka / Event Streaming", "$0", "$55", "$200", "$1,000"],
        ["AI/LLM Tokens (Claude)", "$10", "$100", "$475", "$2,350"],
        ["Email (Resend)", "$0", "$30", "$75", "$300"],
        ["SMS (Twilio)", "$2", "$25", "$90", "$450"],
        ["CDN + Storage (S3/CF)", "$6", "$35", "$170", "$500"],
        ["Monitoring", "$0", "$30", "$80", "$200"],
        ["Search (Elasticsearch)", "$0", "$75", "$300", "$700"],
        ["Domain + SSL", "$15", "$15", "$15", "$15"],
    ],
)

doc.add_paragraph()
doc.add_heading("Annual Compliance & Security Costs", level=2)
add_table(doc,
    ["Cost", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["PCI-DSS Compliance", "$250", "$250", "$500", "$500", "$500"],
        ["Penetration Testing", "$5,000", "$5,000", "$10,000", "$10,000", "$15,000"],
        ["SOC 2 Type II Audit", "$0", "$25,000", "$30,000", "$35,000", "$40,000"],
        ["Apple Developer Program", "$99", "$99", "$99", "$99", "$99"],
    ],
)

# ============================================================
# 12. MARKETING & USER ACQUISITION
# ============================================================
doc.add_heading("12. Marketing & User Acquisition", level=1)

doc.add_heading("Canadian Marketing (With CI)", level=2)
add_table(doc,
    ["Channel", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["CI Referral (partnership)", "$0", "$0", "$0", "$0", "$0"],
        ["Organic Marketing", "$6,000", "$12,000", "$18,000", "$24,000", "$30,000"],
        ["Paid Media", "$1,800", "$4,600", "$53,000", "$100,400", "$148,800"],
        ["Total", "$7,800", "$16,600", "$71,000", "$124,400", "$178,800"],
    ],
)

doc.add_paragraph()
doc.add_heading("Canadian Cost per Visitor", level=3)
add_table(doc,
    ["Channel", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["CI Referral", "$0.00", "$0.00", "$0.00", "$0.00", "$0.00"],
        ["Organic Search", "$0.27", "$0.30", "$0.36", "$0.40", "$0.40"],
        ["Paid Media", "$0.10", "$0.13", "$1.33", "$2.01", "$2.29"],
        ["Blended", "$0.06", "$0.10", "$0.34", "$0.48", "$0.56"],
    ],
)

doc.add_paragraph()
doc.add_heading("US Marketing Spend", level=2)
add_table(doc,
    ["Year", "Organic", "Paid Media", "Total", "Notes"],
    [
        ["FY2027 (8 mo)", "$10,000", "$40,000", "$50,000", "Bootstrap: leverage CA press"],
        ["FY2028", "$36,000", "$150,000", "$186,000", "Brand building, paid acquisition"],
        ["FY2029", "$60,000", "$240,000", "$300,000", "Scale: content + paid at volume"],
        ["FY2030", "$84,000", "$320,000", "$404,000", "Growth: network effects reduce CAC"],
        ["FY2031", "$108,000", "$400,000", "$508,000", "Maturity: organic share increasing"],
    ],
)

doc.add_paragraph()
doc.add_paragraph(
    "US marketing is heavier than CA (no CI traffic equivalent) but lighter than Scenario B "
    "(without CI) at equivalent user counts, because CA credibility reduces cold-market "
    "acquisition costs by an estimated 20\u201330%."
)

# ============================================================
# 13. ORGANIZATIONAL COSTS
# ============================================================
doc.add_heading("13. Organizational Costs — All Four Entities", level=1)

doc.add_heading("13a. Harness Exchange — CA For-Profit (CBCA / Ontario)", level=2)
doc.add_paragraph("One-time formation: $590 CAD")
add_table(doc,
    ["Cost", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["Annual return", "$50", "$50", "$50", "$50", "$50"],
        ["T2 corporate tax return", "$1,200", "$1,500", "$2,000", "$2,500", "$3,000"],
        ["HST/GST filing", "$1,200", "$1,200", "$2,000", "$2,000", "$2,000"],
        ["Bookkeeping", "$3,600", "$6,000", "$9,600", "$12,000", "$18,000"],
        ["Financial review/audit", "$0", "$5,000", "$10,000", "$15,000", "$20,000"],
        ["D&O insurance", "$3,000", "$3,000", "$5,000", "$7,000", "$10,000"],
        ["Commercial general liability", "$800", "$800", "$1,200", "$1,500", "$2,000"],
        ["Cyber / tech E&O insurance", "$2,000", "$3,000", "$5,000", "$8,000", "$12,000"],
        ["Legal retainer", "$5,000", "$8,000", "$12,000", "$15,000", "$20,000"],
        ["Board of Directors", "$0", "$0", "$5,000", "$10,000", "$20,000"],
        ["Registered agent", "$500", "$500", "$500", "$500", "$500"],
        ["Business banking", "$0", "$600", "$1,200", "$1,500", "$1,500"],
        ["Total (annual, CAD)", "$17,350", "$29,650", "$53,550", "$75,050", "$109,050"],
    ],
)

doc.add_paragraph()
doc.add_heading("13b. Harness Good — CA Not-for-Profit (CNCA)", level=2)
doc.add_paragraph("One-time formation: $3,200 CAD")
add_table(doc,
    ["Cost", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["Annual return", "$12", "$12", "$12", "$12", "$12"],
        ["T3010 charity return", "$1,500", "$2,000", "$2,500", "$3,000", "$3,500"],
        ["Bookkeeping", "$2,400", "$3,600", "$6,000", "$9,600", "$12,000"],
        ["Financial audit", "$0", "$8,000", "$10,000", "$15,000", "$20,000"],
        ["General liability insurance", "$500", "$500", "$800", "$1,000", "$1,500"],
        ["Cyber insurance", "$1,500", "$2,000", "$3,000", "$5,000", "$8,000"],
        ["Legal retainer", "$3,000", "$5,000", "$8,000", "$10,000", "$12,000"],
        ["Board of Directors", "$0", "$0", "$0", "$0", "$0"],
        ["Total (annual, CAD)", "$8,912", "$21,112", "$30,312", "$43,612", "$57,012"],
    ],
)

doc.add_paragraph()
doc.add_heading("13c. Harness Exchange US — Delaware C-Corp", level=2)
doc.add_paragraph("One-time formation: $159 USD")
add_table(doc,
    ["Cost", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["Delaware franchise tax", "$400", "$400", "$400", "$400", "$400"],
        ["Delaware annual report", "$50", "$50", "$50", "$50", "$50"],
        ["Registered agent", "$50", "$50", "$50", "$50", "$50"],
        ["Form 1120 tax return", "$1,500", "$2,000", "$3,000", "$3,500", "$5,000"],
        ["Financial audit", "$0", "$0", "$10,000", "$15,000", "$20,000"],
        ["D&O insurance", "$2,000", "$2,000", "$4,000", "$6,000", "$8,000"],
        ["General liability + cyber", "$1,500", "$2,000", "$3,000", "$5,000", "$8,000"],
        ["Legal (US corporate)", "$3,000", "$5,000", "$8,000", "$10,000", "$15,000"],
        ["Board of Directors", "$0", "$0", "$3,000", "$5,000", "$10,000"],
        ["Total (annual, USD)", "$8,500", "$11,500", "$31,500", "$45,000", "$66,500"],
    ],
)

doc.add_paragraph()
doc.add_heading("13d. Harness Good US Foundation — Delaware 501(c)(3)", level=2)
doc.add_paragraph("One-time formation: $3,709 USD")
add_table(doc,
    ["Cost", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["Delaware annual report", "$25", "$25", "$25", "$25", "$25"],
        ["Delaware franchise tax", "$0", "$0", "$0", "$0", "$0"],
        ["Form 990 preparation", "$1,000", "$1,500", "$2,000", "$2,000", "$2,000"],
        ["Financial audit", "$0", "$5,000", "$8,000", "$12,000", "$15,000"],
        ["General liability + cyber", "$1,000", "$1,500", "$2,500", "$4,000", "$6,000"],
        ["Legal (US NFP counsel)", "$2,000", "$3,000", "$5,000", "$8,000", "$10,000"],
        ["Board of Directors", "$0", "$0", "$0", "$0", "$0"],
        ["Total (annual, USD)", "$4,025", "$11,025", "$17,525", "$26,025", "$33,025"],
    ],
)

doc.add_paragraph()
doc.add_heading("US Compliance & Organizational Costs", level=2)
add_table(doc,
    ["Cost", "Y1 (8 mo)", "Y2", "Y3", "Y4", "Y5"],
    [
        ["State charity registrations", "$3,333", "$8,000", "$12,000", "$15,000", "$18,000"],
        ["US legal counsel", "$6,667", "$15,000", "$20,000", "$25,000", "$30,000"],
        ["Form 990 preparation", "$0", "$3,500", "$5,000", "$7,500", "$10,000"],
        ["US bookkeeping", "$2,000", "$6,000", "$9,000", "$12,000", "$15,000"],
        ["US D&O / liability insurance", "$1,333", "$3,000", "$5,000", "$8,000", "$12,000"],
    ],
)

doc.add_paragraph()
doc.add_heading("Consolidated Summary (All Entities)", level=2)
add_table(doc,
    ["Entity", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
    [
        ["Harness Exchange (CAD)", "$17,350", "$29,650", "$53,550", "$75,050", "$109,050"],
        ["Harness Good (CAD)", "$8,912", "$21,112", "$30,312", "$43,612", "$57,012"],
        ["US For-Profit (USD)", "$8,500", "$11,500", "$31,500", "$45,000", "$66,500"],
        ["US Not-for-Profit (USD)", "$4,025", "$11,025", "$17,525", "$26,025", "$33,025"],
    ],
)

doc.add_paragraph()
doc.add_paragraph(
    "Note: US amounts are in USD. At a projected USD/CAD rate of ~1.36, US entities add "
    "approximately CAD $17,000 (Year 1) to CAD $135,000 (Year 5) to consolidated group costs."
)

# ============================================================
# 14. KEY MODELING DECISIONS
# ============================================================
doc.add_heading("14. Key Modeling Decisions", level=1)

add_table(doc,
    ["Decision", "Choice", "Rationale"],
    [
        ["Currency", "CAD (CA entities), USD (US entities)", "Harness Exchange is Canadian (ASPE); US entities report in USD"],
        ["Fiscal year", "Calendar year (Jan\u2013Dec)", "Standard for CA for-profit entities"],
        ["Fee model", "% of donation volume", "Industry standard; aligns incentives with charity partners"],
        ["Fee rate decline", "3.5% \u2192 3.2% over 5 yr", "Volume discount expectation; stays competitive at scale"],
        ["Infrastructure scaling", "Linear interpolation between tiers", "Smooth cost curve; avoids step-function jumps"],
        ["Marketing spend", "Scenario-differentiated", "Core question: what is CI traffic worth in avoided marketing spend?"],
        ["Year 5 convergence", "CA: 55K; US: 200K; combined 255K", "US market ~10\u00d7 larger; bootstrap accelerates growth"],
    ],
)

# ============================================================
# 15. KNOWN GAPS
# ============================================================
doc.add_heading("15. Known Gaps", level=1)

gaps = [
    "Staffing costs: No headcount / salary projections (assumes founder-funded early stage)",
    "Churn: No user or charity partner attrition modeling",
    "Currency risk: CA entities' cloud costs are USD-denominated; no FX hedging modeled",
    "Capital expenditure: All costs modeled as OpEx (SaaS/cloud-native)",
    "Tax implications: No income tax, HST/GST, or SR&ED credit projections",
    "Intercompany transactions: Related party fees, licensing, cost-sharing between the four entities not yet modeled",
]
for g in gaps:
    doc.add_paragraph(g, style="List Bullet")

# ============================================================
# SOURCES
# ============================================================
doc.add_heading("Sources & References", level=1)

sources = [
    "Blackbaud Institute, Charitable Giving Report (2023)",
    "M+R Benchmarks 2024 — mrbenchmarks.com",
    "CanadaHelps Giving Report 2023 — canadahelps.org",
    "Network for Good, Digital Giving Trends (2023)",
    "CRA T1 Individual Tax Statistics, 2023 tax year — canada.ca",
    "Charity Intelligence Web Traffic Analysis, April 2026 (internal)",
    "CI Webstats daily log (server-side, 2009–2025) and GA4 (2023–2025)",
    "Google Search Console data, Dec 2024–Nov 2025 (internal)",
    "Partnerize State of the Nation 2024",
    "Ahrefs, \"How Long Does SEO Take?\" (2023)",
    "First Page Sage, Average CAC by Industry (2024)",
    "First Page Sage, Marketing ROI by Channel (2024)",
    "Uber S-1 Filing (2019); Lyft 10-K (2020)",
    "Crunchbase — GoFundMe funding data",
    "Stripe Canada & US pricing — stripe.com",
    "Anthropic API pricing — anthropic.com",
    "Bloomerang pricing & donor retention reports",
    "Donorbox Pro pricing — donorbox.org",
    "CPA Canada — cpacanada.ca",
    "Corporations Canada — fee schedules",
    "Ontario Business Registry — cost schedules",
    "Delaware Division of Corporations — fee schedule",
    "National Association of State Charity Officials — nasconet.org",
    "PCI Security Standards Council — pcisecuritystandards.org",
    "Vanta SOC 2 Audit Cost Guide — vanta.com",
    "Classy, The State of Modern Philanthropy (2024)",
]
for s in sources:
    doc.add_paragraph(s, style="List Bullet")

# ============================================================
# SAVE
# ============================================================
out_path = "/home/selby/abt/docs/Harness-Give-Forecast-Assumptions-CI-US.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
